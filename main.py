import whisper
from docx import Document
import os

model = whisper.load_model("base.en")

audio_out = input("\n📼 Enter the path of the audio file.mp3: ")

# Remove single and double quotes from the file path
audio_out = audio_out.replace("'", "").replace('"', '')

## check mp3 file
if audio_out[-4:] != ".mp3":
    print("🚫 Please enter a valid mp3 file")
    input("\nPress Enter to exit...")
    exit()

audio_out = os.path.normpath(audio_out)

# Get the file name (without extension) from the audio file path
audio_file_name = os.path.splitext(os.path.basename(audio_out))[0]

# script_dir = os.path.dirname(os.path.abspath(__file__)) + "/models/"
# model_dir = os.path.join(script_dir, "base.en.pt")

# result = model.transcribe(audio_out, verbose=False, fp16=False, download_root=model_dir)
result = model.transcribe(audio_out, verbose=False, fp16=False)
transcription_data = result["segments"]

# Specify the path and name of the Word file
output_word_file = f"{audio_file_name}.docx"

# Create a Word document
doc = Document()

# Write the transcript title in bold
title_paragraph = doc.add_paragraph()
title_run = title_paragraph.add_run(f"Transcript of {audio_file_name}.mp3\n\n")
title_run.bold = True

# Write the transcribed text with timestamps to the Word document
for entry in transcription_data:
    start = entry["start"]
    end = entry["end"]
    text = entry["text"]

    # Format the timestamp as HH:MM:SS
    start_time = f"{int(start) // 3600:02d}:{(int(start) % 3600) // 60:02d}:{int(start) % 60:02d}"
    end_time = f"{int(end) // 3600:02d}:{(int(end) % 3600) // 60:02d}:{int(end) % 60:02d}"
    timestamp = f"{start_time} --> {end_time}"

    # Add timestamp and transcribed text to the document
    doc.add_paragraph(f"[{timestamp}] : {text}\n")

# Save the Word document
doc.save(output_word_file)

print(f"\n🎉 Perfect! Transcribed file has been saved to: {output_word_file}")
